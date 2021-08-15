import datetime
from tkinter import messagebox
from docx import Document
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import os
import comtypes.client
from AdminAPP import PvFile

class PV:
    def word(self, numParcel):
        document = Document()
        table = document.add_table(rows=3, cols=3)

        for section in document.sections:
            section.top_margin = Cm(0.25)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1.5)

        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ancfcc.png"
        pic = table.cell(0, 1)
        pic.width = Cm(9)
        img2 = pic.add_paragraph()
        r = img2.add_run()
        r.add_picture(img, width=Cm(3), height=Cm(2))
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        req = table.cell(1, 1)
        req.text = 'Réquisition n° : ...........'
        req.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        req.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = req.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

        numPage = table.cell(1, 2)
        numPage.text = 'Page n° : '
        numPage.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        numPage.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = numPage.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

        c2 = table.cell(0, 2)
        c2.text = 'IF. 85E bis'
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = c2.paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)

        info_left = table.cell(2, 0)
        info_left.width = Cm(8)
        info_left.text = 'Repérage \n ' \
                         'Mappe n° :........ \n ' \
                         'Coordonnées du Centroïde:  \n ' \
                         'X (m) :....... \n ' \
                         'Y (m) :.......'
        info_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        info_left.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = info_left.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

        info_right = table.cell(2, 2)
        info_right.width = Cm(8)
        info_right.text = 'Nombre de parcelles : Unique \n Parcelle n° :  1 \n Sous-Zone n° : 1'
        info_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        info_right.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = info_right.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

# TABLE2
        r = 6
        table2 = document.add_table(rows=r, cols=4, style='Table Grid')

        riv = table2.cell(0, 0)
        riv.width = Cm(9)
        riv.text = 'RIVERAINS \n Nom des propriétaires riverains, \n avec leur adresse, s’il y a lieu. \n Désignation des tenants et aboutissants'
        riv.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        riv.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = riv.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        indi = table2.cell(0, 1)
        indi.text = "Indication \n sur la \n présence ou l'absence des riverains"
        indi.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        indi.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = indi.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        numBorn = table2.cell(0, 2)
        numBorn.text = "NUMEROS \n DES \n BORNES"
        numBorn.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        numBorn.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = numBorn.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        limit = table2.cell(0, 3)
        limit.width = Cm(15)
        limit.text = "PARTICULARITES DES LIMITES \n Signaler les particularités des limites qui ne ressortiraient pas\n " \
                     "nettement du plan, en particulier la non-conformité avec des \n " \
                     "limites préexistantes. Noter les interruptions et reprise de \n " \
                     "bornages, les interventions dans l’ordre ou elles se produisent."
        limit.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        limit.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = limit.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        try:
            filename = 'PV_P' + numParcel + '.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV', filename)
            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV'))
            document.save(filepath)
            msg = messagebox.showinfo("PV", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            PvFile.PvFile()

        except:
            msg = messagebox.showerror("PV", "Veuillez fermer le fichier")

    def pdf(self, numParcel):
        document = Document()
        table = document.add_table(rows=3, cols=3)

        for section in document.sections:
            section.top_margin = Cm(0.25)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1.5)

        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ancfcc.png"
        pic = table.cell(0, 1)
        pic.width = Cm(9)
        img2 = pic.add_paragraph()
        r = img2.add_run()
        r.add_picture(img, width=Cm(3), height=Cm(2))
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        req = table.cell(1, 1)
        req.text = 'Réquisition n° : ...........'
        req.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        req.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = req.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

        numPage = table.cell(1, 2)
        numPage.text = 'Page n° : '
        numPage.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        numPage.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = numPage.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

        c2 = table.cell(0, 2)
        c2.text = 'IF. 85E bis'
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = c2.paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)

        info_left = table.cell(2, 0)
        info_left.width = Cm(8)
        info_left.text = 'Repérage \n ' \
                         'Mappe n° :........ \n ' \
                         'Coordonnées du Centroïde:  \n ' \
                         'X (m) :....... \n ' \
                         'Y (m) :.......'
        info_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        info_left.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = info_left.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

        info_right = table.cell(2, 2)
        info_right.width = Cm(8)
        info_right.text = 'Nombre de parcelles : Unique \n Parcelle n° :  1 \n Sous-Zone n° : 1'
        info_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        info_right.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = info_right.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

# TABLE2
        r = 6
        table2 = document.add_table(rows=r, cols=4, style='Table Grid')

        riv = table2.cell(0, 0)
        riv.width = Cm(9)
        riv.text = 'RIVERAINS \n Nom des propriétaires riverains, \n avec leur adresse, s’il y a lieu. \n Désignation des tenants et aboutissants'
        riv.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        riv.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = riv.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        indi = table2.cell(0, 1)
        indi.text = "Indication \n sur la \n présence ou l'absence des riverains"
        indi.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        indi.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = indi.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        numBorn = table2.cell(0, 2)
        numBorn.text = "NUMEROS \n DES \n BORNES"
        numBorn.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        numBorn.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = numBorn.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        limit = table2.cell(0, 3)
        limit.width = Cm(15)
        limit.text = "PARTICULARITES DES LIMITES \n Signaler les particularités des limites qui ne ressortiraient pas\n " \
                     "nettement du plan, en particulier la non-conformité avec des \n " \
                     "limites préexistantes. Noter les interruptions et reprise de \n " \
                     "bornages, les interventions dans l’ordre ou elles se produisent."
        limit.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        limit.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = limit.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        try:
            filename = 'PV_P' + numParcel + '.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV', filename)
            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV'))
            document.save(filepath)
            # PDF
            wordFile = 'PV_P' + numParcel + '.docx'
            pdfFile = 'PV_P' + numParcel + '.pdf'
            wdFormatPDF = 17
            in_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV', wordFile)
            out_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV', pdfFile)
            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\PV'))
            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=wdFormatPDF)
            doc.close()
            word.Quit()
            msg = messagebox.showinfo("PV", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            PvFile.PvFile()

        except:
            msg = messagebox.showerror("PV", "Veuillez fermer le fichier")





