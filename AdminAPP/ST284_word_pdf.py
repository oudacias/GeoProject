import datetime
from tkinter import messagebox
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Mm
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.enum.table import WD_TABLE_DIRECTION
import os
import comtypes.client
from AdminAPP import ST284File

class ST284:
    def word(self, nomParcel, requis, titr, numParcel):
        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ST284.png"

        document = Document()
        table = document.add_table(rows=4, cols=3)

        for section in document.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(297)  # for A4 paper
            section.page_height = Mm(210)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        titre = table.cell(1, 0)
        titre.text = 'Agence Nationale \n de la Conservation \n Foncière du Cadastre \n et de la Cartographie '
        titre.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Cambria'
        run.font.size = Pt(13)

        titre2 = table.cell(2, 0)
        titre2.width = Cm(7)
        titre2.text = '---------\n Service \n du Cadastre \n de Tétouan'
        titre2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        titre2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre2.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Cambria'
        run.font.size = Pt(13)

        titre3 = table.cell(1, 1)
        titre3.width = Cm(16)
        titre3.text = 'CALCUL DE \n CONTENANCES'
        titre3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre3.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'calibri'
        run.font.size = Pt(26)

        titre4 = table.cell(2, 1)
        titre4.text = " Propriété dite : " + nomParcel + "\n Nature de l'affaire :  IFE \n Réquisition :" + \
                      requis + '            ' + "Titre :" + titr
        titre4.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre4.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = titre4.paragraphs[0].runs[0]
        run.font.italic = True
        run.font.name = 'calibri'
        run.font.size = Pt(14)

        cel = table.cell(0, 2)
        cel2 = table.cell(2, 2)
        cel3 = cel.merge(cel2)
        cel3.width = Cm(7)

        img2 = cel3.add_paragraph()
        r = img2.add_run()
        r.add_picture(img)
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        c1 = table.cell(3, 0)
        c1.text = "Système : LAMBERT"
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c1.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)

        c2 = table.cell(3, 1)
        c2.text = 'P' + numParcel
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c2.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)

        c3 = table.cell(3, 2)
        c3.text = "Coordonnées: CENTRE"
        c3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = c3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)

# -- //table2/ --
        r = 5
        table2 = document.add_table(rows=r, cols=4, style='Table Grid')

        X = table2.cell(0, 0)
        X.width = Cm(8)
        X.text = 'X'
        run = X.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(15)
        X.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        X.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        born = table2.cell(0, 1)
        born.width = Cm(8)
        born.text = 'Bornes'
        run = born.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(15)
        born.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        born.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        Y = table2.cell(0, 2)
        Y.width = Cm(8)
        Y.text = 'Y'
        run = Y.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(15)
        Y.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        Y.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        ref = table2.cell(0, 3)
        ref.width = Cm(16)
        ref.text = 'Références'
        run = ref.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(15)
        ref.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        ref.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# //table3// --
        table3 = document.add_table(rows=3, cols=1, style='Table Grid')
        val = table3.cell(0, 0)
        val.text = '                                                                        S                                                                         =                              00 Ha 05 A 60.7912 CA \n' \
                   '                                                   CORRECTION LAMBERT                                                      =                             -00 Ha 00 a 00.4183 ca'
        run = val.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        val.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        val2 = table3.cell(1, 0)
        val2.text = '                                                       SURFACE CORRIGEE                                                        =                              00 Ha 05 a 60.3730 ca'
        run = val2.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        val2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        val3 = table3.cell(2, 0)
        val3.text = '                                                   CONTENANCE ADOPTEE                                                     =                               00 Ha 05 a 60 ca'
        run = val3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        val3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        try:
            filename = 'ST284_P' + numParcel + '.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ST284', filename)
            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ST284')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ST284'))
            document.save(filepath)
            msg = messagebox.showinfo("ST284", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            ST284File.ST284()

        except:
            msg = messagebox.showerror("ST284", "Veuillez fermer le fichier")


    def pdf(self, nomParcel, requis, titr, numParcel):
        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ST284.png"

        document = Document()
        table = document.add_table(rows=4, cols=3)

        for section in document.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(297)  # for A4 paper
            section.page_height = Mm(210)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        titre = table.cell(1, 0)
        titre.text = 'Agence Nationale \n de la Conservation \n Foncière du Cadastre \n et de la Cartographie '
        titre.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Cambria'
        run.font.size = Pt(13)

        titre2 = table.cell(2, 0)
        titre2.width = Cm(7)
        titre2.text = '---------\n Service \n du Cadastre \n de Tétouan'
        titre2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        titre2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre2.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Cambria'
        run.font.size = Pt(13)

        titre3 = table.cell(1, 1)
        titre3.width = Cm(16)
        titre3.text = 'CALCUL DE \n CONTENANCES'
        titre3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre3.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'calibri'
        run.font.size = Pt(26)

        titre4 = table.cell(2, 1)
        titre4.text = " Propriété dite : " + nomParcel + "\n Nature de l'affaire :  IFE \n Réquisition :" + \
                      requis + '            ' + "Titre :" + titr
        titre4.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre4.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = titre4.paragraphs[0].runs[0]
        run.font.italic = True
        run.font.name = 'calibri'
        run.font.size = Pt(14)

        cel = table.cell(0, 2)
        cel2 = table.cell(2, 2)
        cel3 = cel.merge(cel2)
        cel3.width = Cm(7)

        img2 = cel3.add_paragraph()
        r = img2.add_run()
        r.add_picture(img)
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        c1 = table.cell(3, 0)
        c1.text = "Système : LAMBERT"
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c1.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)

        c2 = table.cell(3, 1)
        c2.text = 'P' + numParcel
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c2.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)

        c3 = table.cell(3, 2)
        c3.text = "Coordonnées: CENTRE"
        c3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = c3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)

# -- //table2/ --
        r = 5
        table2 = document.add_table(rows=r, cols=4, style='Table Grid')

        X = table2.cell(0, 0)
        X.width = Cm(8)
        X.text = 'X'
        run = X.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(15)
        X.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        X.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        born = table2.cell(0, 1)
        born.width = Cm(8)
        born.text = 'Bornes'
        run = born.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(15)
        born.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        born.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        Y = table2.cell(0, 2)
        Y.width = Cm(8)
        Y.text = 'Y'
        run = Y.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(15)
        Y.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        Y.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        ref = table2.cell(0, 3)
        ref.width = Cm(16)
        ref.text = 'Références'
        run = ref.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(15)
        ref.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        ref.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# //table3// --
        table3 = document.add_table(rows=3, cols=1, style='Table Grid')
        val = table3.cell(0, 0)
        val.text = '                                                                        S                                                                         =                              00 Ha 05 A 60.7912 CA \n' \
                   '                                                   CORRECTION LAMBERT                                                      =                             -00 Ha 00 a 00.4183 ca'
        run = val.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        val.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        val2 = table3.cell(1, 0)
        val2.text = '                                                       SURFACE CORRIGEE                                                        =                              00 Ha 05 a 60.3730 ca'
        run = val2.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        val2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        val3 = table3.cell(2, 0)
        val3.text = '                                                   CONTENANCE ADOPTEE                                                     =                               00 Ha 05 a 60 ca'
        run = val3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        val3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        try:
            filename = 'ST284_P' + numParcel + '.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ST284', filename)
            document.save(filepath)
            # PDF
            wordFile = 'ST284_P' + numParcel + '.docx'
            pdfFile = 'ST284_P' + numParcel + '.pdf'
            wdFormatPDF = 17
            in_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ST284', wordFile)
            out_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ST284', pdfFile)

            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ST284')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ST284'))

            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=wdFormatPDF)
            doc.close()
            word.Quit()
            msg = messagebox.showinfo("ST284", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            ST284File.ST284()

        except:
            msg = messagebox.showerror("ST284", "Veuillez fermer le fichier")





