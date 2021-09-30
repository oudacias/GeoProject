from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Mm
from docx.shared import Inches, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
import os
from tkinter import filedialog

project_directory = os.path.abspath(os.curdir)
img = project_directory + "/icons/ancfcc2.png"

import tkinter as tk
from tkinter import *
from UserAPP import Autocompletecombox
import datetime

class Zn2(tk.Tk):
    def __init__(self, *args, **kwargs):
        tk.Tk.__init__(self, *args, **kwargs)
        self.title("ZN2")
        self.geometry("450x240+430+180")
        self.resizable(0, 0)
        self.config(bg="#F7F9F9")

        frmLab = LabelFrame(self, text="ZN2", font=("Times New Roman", 14), bg="#F7F9F9")
        frmLab.place(x=10, y=5, width=430, height=225)


        btn = Button(frmLab, text="Générer", font=("Times New Roman", 11), bg="#4EB1FA", fg="#FFFFFF", relief=FLAT, activebackground="#4EB1FA", activeforeground="#000000", command=self.generate)
        btn.place(x=270, y=140)

        frm3 = Frame(frmLab, bg="#379BF3", width=134, height=28)
        frm3.place(x=100, y=140)
        filterframe = Button(frm3, text="Ouvrir la pièce générée", bg="#CEE6F3", fg="#379BF3", relief='flat', activebackground="#CEE6F3", activeforeground="#379BF3", command=self.open)
        filterframe.place(x=1, y=1)

    def generate(self):
        document = Document()
        table = document.add_table(rows=3, cols=2)

        a = table.cell(0,0).text = 'ZN2'
        #a.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        b = table.cell(1,0).text= 'Carnet n°: \n Page n°'

        #run = b.paragraphs[0].runs[0]
        #run.font.size = Pt(12)
        #b.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        img2 = table.rows[1].cells[1].add_paragraph()
        r = img2.add_run()
        r.add_picture(img, width=Cm(1.5), height=Cm(1.5))
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        self.filename = 'zn2.docx'

        #filepath=r'C:\Users\Administrator\Desktop\python test\update_test'+filename
        #filepath = r'C:\Users\LAILA\PycharmProjects\files'
        filepath = os.path.join(r'C:\Users\LAILA\PycharmProjects\files', self.filename)
        document.save(filepath)

    def open(self):
        path = r'C:\Users\LAILA\PycharmProjects\files'
        options = {
                    'initialdir': path,
                    'title': 'Choose your file', }
        filedialog.askopenfile(**options)

        print('open')

if __name__ == "__main__":
    app = Zn2()
    app.mainloop()