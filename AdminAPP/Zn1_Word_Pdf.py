import datetime
from tkinter import messagebox
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Mm
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.enum.table import WD_TABLE_DIRECTION
import os
import comtypes.client
from AdminAPP import Zn1File

class ZN1:
    def word(self, nomFr, prenomFr, prov, sZone, cercle, commun, adress, cin, tel, douar, numParcel, nomParcel, oppos, const, specul, tpSol, surf_Ha):
        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ancfcc2.png"

        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.rtl = True
        font.name = 'Calibri'
        font.size = Pt(11)

        # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        for section in document.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(297)  # for A4 paper
            section.page_height = Mm(210)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        tab = document.add_table(rows=1, cols=1)

        '''title = tab.cell(0, 8)
        title.text = 'ZN1'
        title.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = title.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(13)'''

        img2 = tab.rows[0].cells[0].add_paragraph()
        r = img2.add_run()
        r.add_picture(img, width=Cm(1.5), height=Cm(1.5))
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        table = document.add_table(rows=10, cols=18, style='Table Grid')

        a = table.cell(0, 0)
        b = table.cell(1, 4)
        A = a.merge(b)
        A.text = 'Renseignements donnés par : ' + nomFr + ' ' + prenomFr
        A.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        run = A.paragraphs[0].runs[0]
        run.font.bold = True

        a2 = table.cell(0, 5)
        b2 = table.cell(1, 14)
        livret = a2.merge(b2)
        livret.text = ' Livret remis à : ' + nomFr + ' ' + prenomFr
        livret.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        run = livret.paragraphs[0].runs[0]
        run.font.bold = True

# -- IMAGE --
        pic = table.cell(0, 15)
        pic2 = table.cell(1, 17)
        pic3 = pic.merge(pic2)

        pic3.text = 'ZN1'
        run = pic3.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(14)
        pic3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        paragph = pic3.add_paragraph('Carnet n°: ' + '...' + '\n Page n°' + '...')
        paragph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

        '''img2 = table.rows[1].cells[15].add_paragraph()
        r = img2.add_run()
        r.add_picture(img, width=Cm(1.5), height=Cm(1.5))
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP '''

# -- PROVINCE --
        e = table.cell(2, 0)
        f = table.cell(2, 1)
        C = e.merge(f)
        C.text = 'Province de'
        C.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = C.paragraphs[0].runs[0]
        run.font.bold = True

        g = table.cell(2, 2)
        h = table.cell(2, 4)
        D = g.merge(h)
        D.text = prov
        run = D.paragraphs[0].runs[0]
        run.font.name = 'Cambria'

        e = table.cell(2, 5)
        f = table.cell(2, 13)
        j = e.merge(f)
        j.text = 'RECONNAISSANCE PARCELLAIRE'
        run = j.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(15)
        j.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        j.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        h = table.cell(2, 14)
        i = table.cell(2, 17)
        k = h.merge(i)
        k.text = 'Sous-Secteur (Zone) n°:' + sZone
        k.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        c1 = table.cell(3, 0)
        c2 = table.cell(3, 1)
        c3 = c1.merge(c2)
        c3.text = 'Cercle de'
        c3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = c3.paragraphs[0].runs[0]
        run.font.bold = True

        c4 = table.cell(3, 2)
        c5 = table.cell(3, 4)
        c6 = c4.merge(c5)
        c6.text = cercle
        run = c6.paragraphs[0].runs[0]
        run.font.name = 'Cambria'

        c7 = table.cell(3, 5)
        c8 = table.cell(3, 13)
        c9 = c7.merge(c8)
        c9.text = 'Commune Rurale : ' + commun
        c9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c9.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(15)

        c10 = table.cell(3, 14)
        c11 = table.cell(3, 17)
        c12 = c10.merge(c11)
        c12.text = 'Date :' + str(datetime.datetime.now())
        c12.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# -- NOM ET ADRESSE DES PROPRIETAIRES --
        x1 = table.cell(4, 0)
        x2 = table.cell(8, 4)
        x3 = x1.merge(x2)
        x3.text = 'NOM ET ADRESSE DES \n PROPRIETAIRES \n (1)'
        x3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        x3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = x3.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(12)
        run.underline = True

        nom = table.cell(9, 0)
        nom2 = table.cell(9, 4)
        nom3 = nom.merge(nom2)
        nom3.text = nomFr + ' ' + prenomFr + '\n' + adress + '\n CIN' + ' ' + cin + '\n TEL' + ' ' + tel
        nom3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        nom3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = nom3.paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)

# -- RENSEIGNEMENTS SUR LES PARCELLES --
        x4 = table.cell(4, 5)
        x5 = table.cell(4, 17)
        x6 = x4.merge(x5)
        x6.text = 'RENSEIGNEMENTS SUR LES PARCELLES'
        x6.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = x6.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(14)

# -- DOUAR --
        x7 = table.cell(5, 5)
        x8 = table.cell(7, 5)
        x9 = x7.merge(x8)
        x9.text = 'Douar'
        x9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        x9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        z10 = table.cell(8, 5)
        z10.text = '(2)'
        z10.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z10.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        drVal = table.cell(9, 5)
        drVal.text = douar
        drVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        drVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = drVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- NUM_PARCELLE --
        x10 = table.cell(5, 6)
        x11 = table.cell(7, 6)
        x12 = x10.merge(x11)
        x12.text = 'N° Plle \n / \n N° Réq'
        x12.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        x12.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z11 = table.cell(8, 6)
        z11.text = '(3)'
        z11.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z11.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        numPolyg = table.cell(9, 6)
        numPolyg.text = numParcel
        numPolyg.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        numPolyg.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = numPolyg.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- PARCELLE --
        y = table.cell(5, 7)
        y1 = table.cell(7, 8)
        y2 = y.merge(y1)
        y2.text = 'Nom \n de la parcelle'
        y2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z18 = table.cell(8, 7)
        z19 = table.cell(8, 8)
        z20 = z18.merge(z19)
        z20.text = '(4)'
        z20.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z20.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        nom_polyg = table.cell(9, 7)
        nom_polyg_2 = table.cell(9, 8)
        nom_polyg_val = nom_polyg.merge(nom_polyg_2)
        nom_polyg_val.text = nomParcel
        nom_polyg_val.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nom_polyg_val.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = nom_polyg_val.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- OPPOSITION --
        y3 = table.cell(5, 9)
        y4 = table.cell(7, 9)
        y5 = y3.merge(y4)
        y5.text = 'Oppo\nsition \n O/N'
        y5.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y5.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z13 = table.cell(8, 9)
        z13.text = '(5)'
        z13.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z13.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        oppoVal = table.cell(9, 9)

        if int(oppos) > 0:
            oppoVal.text = 'O'  # OPPOSITION
        elif int(oppos) == 0:
            oppoVal.text = 'N'  # Non OPPOSITION
        oppoVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        oppoVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = oppoVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- CONSISTANCE --
        y6 = table.cell(5, 10)
        y7 = table.cell(7, 10)
        y8 = y6.merge(y7)
        y8.text = 'NATURE \n TC \n TP \n TB'
        y8.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y8.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z14 = table.cell(8, 10)
        z14.text = '(6)'
        z14.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z14.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        constVal = table.cell(9, 10)
        constVal.text = const
        constVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        constVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = constVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- SPÉCULATION --
        y9 = table.cell(5, 11)
        y10 = table.cell(7, 11)
        y11 = y9.merge(y10)
        y11.text = 'Type de Spécula\ntion TR-IN-MA'
        y11.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y11.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z15 = table.cell(8, 11)
        z15.text = '(7)'
        z15.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z15.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        speclVal = table.cell(9, 11)
        speclVal.text = specul
        speclVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        speclVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = speclVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- TYPE DE SOL --
        y12 = table.cell(5, 12)
        y13 = table.cell(7, 12)
        y14 = y12.merge(y13)
        y14.text = 'Type de Sol \n RME \n HAM \n TIR-DEH \n BIA'
        y14.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y14.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z16 = table.cell(8, 12)
        z16.text = '(8)'
        z16.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z16.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        solVal = table.cell(9, 12)
        solVal.text = tpSol
        solVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        solVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = solVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- SUPÉRFICIE --
        y15 = table.cell(5, 13)
        y16 = table.cell(7, 13)
        y17 = y15.merge(y16)
        y17.text = 'Superficie \n Déclarée (Ha)'
        y17.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y17.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z17 = table.cell(8, 13)
        z17.text = '(9)'
        z17.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z17.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        surfce = table.cell(9, 13)
        surfce.text = surf_Ha
        surfce.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        surfce.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = surfce.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- SITUATION GÉOGRAPHIQUE --
        z7 = table.cell(5, 14)
        z8 = table.cell(7, 17)
        z9 = z7.merge(z8)
        z9.text = 'Situation Géographique \n Mappe : XX-XX-XX / X (m)/ Y(m) \n Observations'
        z9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        a1 = table.cell(8, 14)
        a2 = table.cell(8, 17)
        a3 = a1.merge(a2)
        a3.text = '(10)'
        a3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        a3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        mapp_x_y = table.cell(9, 14)
        mapp_x_y_2 = table.cell(9, 17)
        mapp_x_y_3 = mapp_x_y.merge(mapp_x_y_2)
        mapp_x_y_3.text = '..-..-../.../...'
        mapp_x_y_3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        mapp_x_y_3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = mapp_x_y_3.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- \\ INFORMATION BAS DE PAGE // --
        paragph1 = document.add_paragraph(" (1) A compléter par le numéro CINE ( La Carte Nationale d'Identité Électronique) et numéro de GSM le cas échéant")
        paragph1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph1.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph1.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        font = run.font
        font.color.rgb = RGBColor(0x96, 0x96, 0x96)

        paragph2 = document.add_paragraph("(6) Consistance Matérielle : TC ( Terrain de Culture) –TP (Terrain Implanté) – TB (Terrain Bâti) \n A compléter en Arabe par :")
        paragph2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph2.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph2.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        font = run.font
        font.color.rgb = RGBColor(0x96, 0x96, 0x96)

        table2 = document.add_table(rows=1, cols=12, style='Table Grid')
        table2.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table2.alignment = WD_TABLE_DIRECTION.RTL

        table2.cell(0, 0).text = u'أرض فلاحية'
        table2.cell(0, 1).text = u'ارض بورية'
        table2.cell(0, 2).text = u'ارض زراعية مسقية'
        table2.cell(0, 3).text = u'ارض مغروسة'
        table2.cell(0, 4).text = u'ارض مسقية مغروسة'
        table2.cell(0, 5).text = u'ارض عارية'
        table2.cell(0, 6).text = u'أحراش'
        table2.cell(0, 7).text = u'أرض بها بناية'
        table2.cell(0, 8).text = u'ارض بها بنايات'
        table2.cell(0, 9).text = u'مقبرة'
        table2.cell(0, 10).text = u'مسجد'
        table2.cell(0, 11).text = u'مركب عقاري البيعة  -  مرفق رياضي'

        for column_index in range(len(table2.columns)):
            table2.cell(0, column_index).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table2.cell(0, column_index).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            column_index += 1

        for col_index in range(len(table2.columns)):
            if col_index == 11:
                table2.cell(0, col_index).width = Cm(3.5)
            else:
                table2.cell(0, col_index).width = Cm(2)
            col_index += 1

        paragph3 = document.add_paragraph("(7) Type de Spéculation : TR ( Culture Traditionnelle)- IN ( Culture Industrielle)- MA (Culture Maraichère-potagère)")
        paragph3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph3.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format = paragph3.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph3.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        font = run.font
        font.color.rgb = RGBColor(0x96, 0x96, 0x96)

        paragph4 = document.add_paragraph("(8) Type de sol : RME (R’mel)- HAM (Hamri)- TIR (Tirs)- DEH (Dehs)- BIA (Beida).")
        paragph4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph4.paragraph_format.line_spacing = 1.0
        paragraph_format = paragph4.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph4.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        font = run.font
        font.color.rgb = RGBColor(0x96, 0x96, 0x96)

        paragph5 = document.add_paragraph("N.B : Les carnets ZN1 doivent être établis sur calque stable ordinaire non végétal permettant leur conservation durable.")
        paragph5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph5.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph5.runs[0]
        run.font.bold = True
        run.font.size = Pt(13)
        run.underline = True
        font = run.font
        font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        # paragraph_format.line_spacing = Pt(5)
        try:
            filename = 'P' + numParcel + '_ZN1.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1', filename)
            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1'))

            document.save(filepath)
            msg = messagebox.showinfo("ZN1", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            Zn1File.Zn1File()

        except:
            msg = messagebox.showerror("ZN1", "Veuillez fermer le fichier")


    def wdFormatPDF(self, nomFr, prenomFr, prov, sZone, cercle, commun, adress, cin, tel, douar, numParcel, nomParcel, oppos, const, specul, tpSol, surf_Ha):
        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ancfcc2.png"

        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        for section in document.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(297)  # for A4 paper
            section.page_height = Mm(210)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        tab = document.add_table(rows=1, cols=1)

        '''title = tab.cell(0, 8)
        title.text = 'ZN1'
        title.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = title.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(13)'''

        img2 = tab.rows[0].cells[0].add_paragraph()
        r = img2.add_run()
        r.add_picture(img, width=Cm(1.5), height=Cm(1.5))
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        table = document.add_table(rows=10, cols=18, style='Table Grid')

        a = table.cell(0, 0)
        b = table.cell(1, 4)
        A = a.merge(b)
        A.text = 'Renseignements donnés par : ' + nomFr + ' ' + prenomFr
        A.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        run = A.paragraphs[0].runs[0]
        run.font.bold = True

        a2 = table.cell(0, 5)
        b2 = table.cell(1, 14)
        livret = a2.merge(b2)
        livret.text = ' Livret remis à : ' + nomFr + ' ' + prenomFr
        livret.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        run = livret.paragraphs[0].runs[0]
        run.font.bold = True

# -- IMAGE --
        pic = table.cell(0, 15)
        pic2 = table.cell(1, 17)
        pic3 = pic.merge(pic2)

        pic3.text = 'ZN1'
        run = pic3.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(14)
        pic3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        paragph = pic3.add_paragraph('Carnet n°: ' + '...' + '\n Page n°' + '...')
        paragph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

        '''img2 = table.rows[1].cells[15].add_paragraph()
        r = img2.add_run()
        r.add_picture(img, width=Cm(1.5), height=Cm(1.5))
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP '''

# -- PROVINCE --
        e = table.cell(2, 0)
        f = table.cell(2, 1)
        C = e.merge(f)
        C.text = 'Province de'
        C.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = C.paragraphs[0].runs[0]
        run.font.bold = True

        g = table.cell(2, 2)
        h = table.cell(2, 4)
        D = g.merge(h)
        D.text =  prov
        run = D.paragraphs[0].runs[0]
        run.font.name = 'Cambria'

        e = table.cell(2, 5)
        f = table.cell(2, 13)
        j = e.merge(f)
        j.text = 'RECONNAISSANCE PARCELLAIRE'
        run = j.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(15)
        j.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        j.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        h = table.cell(2, 14)
        i = table.cell(2, 17)
        k = h.merge(i)
        k.text = 'Sous-Secteur (Zone) n°:' + sZone
        k.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        c1 = table.cell(3, 0)
        c2 = table.cell(3, 1)
        c3 = c1.merge(c2)
        c3.text = 'Cercle de'
        c3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = c3.paragraphs[0].runs[0]
        run.font.bold = True

        c4 = table.cell(3, 2)
        c5 = table.cell(3, 4)
        c6 = c4.merge(c5)
        c6.text = cercle
        run = c6.paragraphs[0].runs[0]
        run.font.name = 'Cambria'

        c7 = table.cell(3, 5)
        c8 = table.cell(3, 13)
        c9 = c7.merge(c8)
        c9.text = 'Commune Rurale : ' + commun
        c9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c9.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(15)

        c10 = table.cell(3, 14)
        c11 = table.cell(3, 17)
        c12 = c10.merge(c11)
        c12.text = 'Date :' + str(datetime.datetime.now())
        c12.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# -- NOM ET ADRESSE DES PROPRIETAIRES --
        x1 = table.cell(4, 0)
        x2 = table.cell(8, 4)
        x3 = x1.merge(x2)
        x3.text = 'NOM ET ADRESSE DES \n PROPRIETAIRES \n (1)'
        x3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        x3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = x3.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(12)
        run.underline = True

        nom = table.cell(9, 0)
        nom2 = table.cell(9, 4)
        nom3 = nom.merge(nom2)
        nom3.text = nomFr + ' ' + prenomFr + '\n' + adress + '\n CIN' + ' ' + cin + '\n TEL' + ' ' + tel
        nom3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        nom3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = nom3.paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)

# -- RENSEIGNEMENTS SUR LES PARCELLES --
        x4 = table.cell(4, 5)
        x5 = table.cell(4, 17)
        x6 = x4.merge(x5)
        x6.text = 'RENSEIGNEMENTS SUR LES PARCELLES'
        x6.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = x6.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(14)

# -- DOUAR --
        x7 = table.cell(5, 5)
        x8 = table.cell(7, 5)
        x9 = x7.merge(x8)
        x9.text = 'Douar'
        x9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        x9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        z10 = table.cell(8, 5)
        z10.text = '(2)'
        z10.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z10.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        drVal = table.cell(9, 5)
        drVal.text = douar
        drVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        drVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = drVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- NUM_PARCELLE --
        x10 = table.cell(5, 6)
        x11 = table.cell(7, 6)
        x12 = x10.merge(x11)
        x12.text = 'N° Plle \n / \n N° Réq'
        x12.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        x12.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z11 = table.cell(8, 6)
        z11.text = '(3)'
        z11.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z11.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        numPolyg = table.cell(9, 6)
        numPolyg.text = numParcel
        numPolyg.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        numPolyg.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = numPolyg.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- PARCELLE --
        y = table.cell(5, 7)
        y1 = table.cell(7, 8)
        y2 = y.merge(y1)
        y2.text = 'Nom \n de la parcelle'
        y2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z18 = table.cell(8, 7)
        z19 = table.cell(8, 8)
        z20 = z18.merge(z19)
        z20.text = '(4)'
        z20.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z20.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        nom_polyg = table.cell(9, 7)
        nom_polyg_2 = table.cell(9, 8)
        nom_polyg_val = nom_polyg.merge(nom_polyg_2)
        nom_polyg_val.text = nomParcel
        nom_polyg_val.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nom_polyg_val.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = nom_polyg_val.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- OPPOSITION --
        y3 = table.cell(5, 9)
        y4 = table.cell(7, 9)
        y5 = y3.merge(y4)
        y5.text = 'Oppo\nsition \n O/N'
        y5.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y5.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z13 = table.cell(8, 9)
        z13.text = '(5)'
        z13.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z13.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        oppoVal = table.cell(9, 9)

        if int(oppos) > 0:
            oppoVal.text = 'O'  # OPPOSITION
        elif int(oppos) == 0:
            oppoVal.text = 'N'  # Non OPPOSITION
        oppoVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        oppoVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = oppoVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- CONSISTANCE --
        y6 = table.cell(5, 10)
        y7 = table.cell(7, 10)
        y8 = y6.merge(y7)
        y8.text = 'NATURE \n TC \n TP \n TB'
        y8.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y8.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z14 = table.cell(8, 10)
        z14.text = '(6)'
        z14.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z14.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        constVal = table.cell(9, 10)
        constVal.text = const
        constVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        constVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = constVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- SPÉCULATION --
        y9 = table.cell(5, 11)
        y10 = table.cell(7, 11)
        y11 = y9.merge(y10)
        y11.text = 'Type de Spécula\ntion TR-IN-MA'
        y11.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y11.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z15 = table.cell(8, 11)
        z15.text = '(7)'
        z15.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z15.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        speclVal = table.cell(9, 11)
        speclVal.text = specul
        speclVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        speclVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = speclVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- TYPE DE SOL --
        y12 = table.cell(5, 12)
        y13 = table.cell(7, 12)
        y14 = y12.merge(y13)
        y14.text = 'Type de Sol \n RME \n HAM \n TIR-DEH \n BIA'
        y14.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y14.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z16 = table.cell(8, 12)
        z16.text = '(8)'
        z16.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z16.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        solVal = table.cell(9, 12)
        solVal.text = tpSol
        solVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        solVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = solVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- SUPÉRFICIE --
        y15 = table.cell(5, 13)
        y16 = table.cell(7, 13)
        y17 = y15.merge(y16)
        y17.text = 'Superficie \n Déclarée (Ha)'
        y17.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y17.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z17 = table.cell(8, 13)
        z17.text = '(9)'
        z17.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z17.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        surfce = table.cell(9, 13)
        surfce.text = surf_Ha
        surfce.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        surfce.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = surfce.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- SITUATION GÉOGRAPHIQUE --
        z7 = table.cell(5, 14)
        z8 = table.cell(7, 17)
        z9 = z7.merge(z8)
        z9.text = 'Situation Géographique \n Mappe : XX-XX-XX / X (m)/ Y(m) \n Observations'
        z9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        a1 = table.cell(8, 14)
        a2 = table.cell(8, 17)
        a3 = a1.merge(a2)
        a3.text = '(10)'
        a3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        a3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        mapp_x_y = table.cell(9, 14)
        mapp_x_y_2 = table.cell(9, 17)
        mapp_x_y_3 = mapp_x_y.merge(mapp_x_y_2)
        mapp_x_y_3.text = '..-..-../.../...'
        mapp_x_y_3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        mapp_x_y_3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = mapp_x_y_3.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- \\ INFORMATION BAS DE PAGE // --
        paragph1 = document.add_paragraph(" (1) A compléter par le numéro CINE ( La Carte Nationale d'Identité Électronique) et numéro de GSM le cas échéant")
        paragph1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph1.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph1.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        font = run.font
        font.color.rgb = RGBColor(0x96, 0x96, 0x96)

        paragph2 = document.add_paragraph("(6) Consistance Matérielle : TC ( Terrain de Culture) –TP (Terrain Implanté) – TB (Terrain Bâti) \n A compléter en Arabe par :")
        paragph2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph2.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph2.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        font = run.font
        font.color.rgb = RGBColor(0x96, 0x96, 0x96)

        table2 = document.add_table(rows=1, cols=12, style='Table Grid')
        table2.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table2.alignment = WD_TABLE_DIRECTION.RTL

        table2.cell(0, 0).text =  u'أرض فلاحية'
        table2.cell(0, 1).text = u'ارض بورية'
        table2.cell(0, 2).text = u'ارض زراعية مسقية'
        table2.cell(0, 3).text = u'ارض مغروسة'
        table2.cell(0, 4).text = u'ارض مسقية مغروسة'
        table2.cell(0, 5).text = u'ارض عارية'
        table2.cell(0, 6).text = u'أحراش'
        table2.cell(0, 7).text = u'أرض بها بناية'
        table2.cell(0, 8).text = u'ارض بها بنايات'
        table2.cell(0, 9).text = u'مقبرة'
        table2.cell(0, 10).text = u'مسجد'
        table2.cell(0, 11).text = u'مركب عقاري البيعة  -  مرفق رياضي'

        for column_index in range(len(table2.columns)):
            table2.cell(0, column_index).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table2.cell(0, column_index).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            column_index += 1

        for col_index in range(len(table2.columns)):
            if col_index == 11:
                table2.cell(0, col_index).width = Cm(3.5)
            else:
                table2.cell(0, col_index).width = Cm(2)
            col_index += 1

        paragph3 = document.add_paragraph("(7) Type de Spéculation : TR ( Culture Traditionnelle)- IN ( Culture Industrielle)- MA (Culture Maraichère-potagère)")
        paragph3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph3.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format = paragph3.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph3.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        font = run.font
        font.color.rgb = RGBColor(0x96, 0x96, 0x96)

        paragph4 = document.add_paragraph("(8) Type de sol : RME (R’mel)- HAM (Hamri)- TIR (Tirs)- DEH (Dehs)- BIA (Beida).")
        paragph4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph4.paragraph_format.line_spacing = 1.0
        paragraph_format = paragph4.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph4.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        font = run.font
        font.color.rgb = RGBColor(0x96, 0x96, 0x96)

        paragph5 = document.add_paragraph("N.B : Les carnets ZN1 doivent être établis sur calque stable ordinaire non végétal permettant leur conservation durable.")
        paragph5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph5.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        run = paragph5.runs[0]
        run.font.bold = True
        run.font.size = Pt(13)
        run.underline = True
        font = run.font
        font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        # paragraph_format.line_spacing = Pt(5)
        try:
            filename = 'P' + numParcel + '_ZN1.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1', filename)
            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1'))
            document.save(filepath)
            print(filepath)
            # PDF
            wordFile = 'P' + numParcel + '_ZN1.docx'
            pdfFile = 'P' + numParcel + '_ZN1.pdf'
            wdFormatPDF = 17
            in_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1', wordFile)
            out_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1', pdfFile)

            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN1'))

            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=wdFormatPDF)
            doc.close()
            word.Quit()
            print('pdf')
            msg = messagebox.showinfo("ZN1", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            Zn1File.Zn1File()

        except:
            msg = messagebox.showerror("ZN1", "Veuillez fermer le fichier")





