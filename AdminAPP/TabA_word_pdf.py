import datetime
from tkinter import messagebox
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Mm
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.enum.table import WD_TABLE_DIRECTION
import os
import comtypes.client
from AdminAPP import TabAFile

class TabA:

    def word(self, nomParcel, requis, titr, numParcel):
        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ST295.png"

        document = Document()
        table = document.add_table(rows=4, cols=3)

        for section in document.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(297)  # for A4 paper
            section.page_height = Mm(210)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        titre = table.cell(1, 0)
        titre.text = 'Agence Nationale \n de la Conservation \n Foncière du Cadastre \n et de la Cartographie '
        titre.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Cambria'
        run.font.size = Pt(14)

        titre2 = table.cell(2, 0)
        titre2.width = Cm(7)
        titre2.text = '---------\n Service \n du Cadastre \n de Tétouan'
        titre2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        titre2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre2.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Cambria'
        run.font.size = Pt(14)

        titre3 = table.cell(1, 1)
        titre3.width = Cm(16)
        titre3.text = 'Tableau A \n Des Contenances'
        titre3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre3.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'calibri'
        run.font.size = Pt(26)

        titre4 = table.cell(2, 1)
        titre4.text = " Propriété dite : " + nomParcel + "\n Nature de l'affaire :  IFE\n Réquisition :" + \
                      requis + '            ' + "Titre :" + titr
        titre4.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre4.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = titre4.paragraphs[0].runs[0]
        run.font.italic = True
        run.font.name = 'calibri'
        run.font.size = Pt(15)

        cel = table.cell(0, 2)
        cel2 = table.cell(2, 2)
        cel3 = cel.merge(cel2)
        cel3.width = Cm(7)

        img2 = cel3.add_paragraph()
        r = img2.add_run()
        r.add_picture(img)
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        c2 = table.cell(3, 1)
        c2.text = 'P' + numParcel
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c2.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)

        table2 = document.add_table(rows=8, cols=8, style='Table Grid')

        c = table2.cell(0, 0)
        c2 = table2.cell(6, 0)
        c3 = c2.merge(c)
        c3.text = 'NATURE \n des travaux \n effectués'
        c3.width = Cm(6)
        c3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        c4 = table2.cell(0, 1)
        c5 = table2.cell(6, 1)
        c6 = c4.merge(c5)
        c6.text = 'PARCELLES'
        c6.width = Cm(4)
        c6.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c6.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c6.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        var = table2.cell(0, 2)
        var2 = table2.cell(6, 2)
        var3 = var.merge(var2)
        var3.text = 'CONTENANCE \n Analytique'
        var3.width = Cm(4)
        var3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        var3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = var3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        var4 = table2.cell(0, 3)
        var5 = table2.cell(6, 3)
        var6 = var4.merge(var5)
        var6.text = 'CONTENANCE \n Graphique \n Vérification'
        var6.width = Cm(4)
        var6.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        var6.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = var6.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        dist = table2.cell(0, 4)
        dist2 = table2.cell(6, 4)
        dist3 = dist.merge(dist2)
        dist3.text = 'DISCORDANCE'
        dist3.width = Cm(2)
        dist3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        dist3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = dist3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        tolr = table2.cell(0, 5)
        tolr2 = table2.cell(6, 5)
        tolr3 = tolr.merge(tolr2)
        tolr3.text = 'TOLERANCE'
        tolr3.width = Cm(2)
        tolr3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tolr3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = tolr3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        cont = table2.cell(0, 6)
        cont2 = table2.cell(6, 6)
        cont3 = cont.merge(cont2)
        cont3.text = 'CONTENANCE \n ADOPTEE'
        cont3.width = Cm(4)
        cont3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cont3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cont3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        obsrv = table2.cell(0, 7)
        obsrv2 = table2.cell(6, 7)
        obsrv3 = obsrv.merge(obsrv2)
        obsrv3.text = 'OBSERVATIONS'
        obsrv3.width = Cm(4)
        obsrv3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        obsrv3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = obsrv3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        try:
            filename = 'ST295_P' + numParcel + '.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\TAB_A', filename)

            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\TAB_A')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\TAB_A'))

            document.save(filepath)
            msg = messagebox.showinfo("ST295", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            TabAFile.Tab_A()
        except:
            msg = messagebox.showerror("ST295", "Veuillez fermer le fichier")


    def pdf(self, nomParcel, requis, titr, numParcel):
        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ST295.png"

        document = Document()
        table = document.add_table(rows=4, cols=3)

        for section in document.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(297)  # for A4 paper
            section.page_height = Mm(210)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        titre = table.cell(1, 0)
        titre.text = 'Agence Nationale \n de la Conservation \n Foncière du Cadastre \n et de la Cartographie '
        titre.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Cambria'
        run.font.size = Pt(14)

        titre2 = table.cell(2, 0)
        titre2.width = Cm(7)
        titre2.text = '---------\n Service \n du Cadastre \n de Tétouan'
        titre2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        titre2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre2.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Cambria'
        run.font.size = Pt(14)

        titre3 = table.cell(1, 1)
        titre3.width = Cm(16)
        titre3.text = 'Tableau A \n Des Contenances'
        titre3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titre3.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'calibri'
        run.font.size = Pt(26)

        titre4 = table.cell(2, 1)
        titre4.text = " Propriété dite : " + nomParcel + "\n Nature de l'affaire :  IFE\n Réquisition :" + \
                      requis + '            ' + "Titre :" + titr
        titre4.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        titre4.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = titre4.paragraphs[0].runs[0]
        run.font.italic = True
        run.font.name = 'calibri'
        run.font.size = Pt(15)

        cel = table.cell(0, 2)
        cel2 = table.cell(2, 2)
        cel3 = cel.merge(cel2)
        cel3.width = Cm(7)

        img2 = cel3.add_paragraph()
        r = img2.add_run()
        r.add_picture(img)
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        c2 = table.cell(3, 1)
        c2.text = 'P' + numParcel
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c2.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)

        table2 = document.add_table(rows=8, cols=8, style='Table Grid')

        c = table2.cell(0, 0)
        c2 = table2.cell(6, 0)
        c3 = c2.merge(c)
        c3.text = 'NATURE \n des travaux \n effectués'
        c3.width = Cm(6)
        c3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        c4 = table2.cell(0, 1)
        c5 = table2.cell(6, 1)
        c6 = c4.merge(c5)
        c6.text = 'PARCELLES'
        c6.width = Cm(4)
        c6.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c6.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c6.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        var = table2.cell(0, 2)
        var2 = table2.cell(6, 2)
        var3 = var.merge(var2)
        var3.text = 'CONTENANCE \n Analytique'
        var3.width = Cm(4)
        var3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        var3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = var3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        var4 = table2.cell(0, 3)
        var5 = table2.cell(6, 3)
        var6 = var4.merge(var5)
        var6.text = 'CONTENANCE \n Graphique \n Vérification'
        var6.width = Cm(4)
        var6.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        var6.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = var6.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        dist = table2.cell(0, 4)
        dist2 = table2.cell(6, 4)
        dist3 = dist.merge(dist2)
        dist3.text = 'DISCORDANCE'
        dist3.width = Cm(2)
        dist3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        dist3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = dist3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        tolr = table2.cell(0, 5)
        tolr2 = table2.cell(6, 5)
        tolr3 = tolr.merge(tolr2)
        tolr3.text = 'TOLERANCE'
        tolr3.width = Cm(2)
        tolr3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tolr3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = tolr3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        cont = table2.cell(0, 6)
        cont2 = table2.cell(6, 6)
        cont3 = cont.merge(cont2)
        cont3.text = 'CONTENANCE \n ADOPTEE'
        cont3.width = Cm(4)
        cont3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cont3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cont3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True

        obsrv = table2.cell(0, 7)
        obsrv2 = table2.cell(6, 7)
        obsrv3 = obsrv.merge(obsrv2)
        obsrv3.text = 'OBSERVATIONS'
        obsrv3.width = Cm(4)
        obsrv3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        obsrv3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = obsrv3.paragraphs[0].runs[0]
        run.font.name = 'calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        try:
            filename = 'ST295_P' + numParcel + '.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\TAB_A', filename)
            document.save(filepath)
            # PDF
            wordFile = 'ST295_P' + numParcel + '.docx'
            pdfFile = 'ST295_P' + numParcel + '.pdf'
            wdFormatPDF = 17
            in_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\TAB_A', wordFile)
            out_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\TAB_A', pdfFile)

            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\TAB_A')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\TAB_A'))

            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=wdFormatPDF)
            doc.close()
            word.Quit()
            msg = messagebox.showinfo("ST284", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            TabAFile.Tab_A()

        except:
            msg = messagebox.showerror("ST284", "Veuillez fermer le fichier")





