import datetime
from tkinter import messagebox
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Mm
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.enum.table import WD_TABLE_DIRECTION
import os
import comtypes.client
from AdminAPP import Zn3File
from pyarabic.unshape import unshaping_word
class ZN3:
    def word(self, nomFr, prenomFr, prov, sZone, cercle, commun, adress, cin, tel, douar, numParcel, nomParcel, motifOppos):
        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ancfcc2.png"

        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.rtl = True
        # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        for section in document.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(297)  # for A4 paper
            section.page_height = Mm(210)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

# -- IMAGE --
        tab = document.add_table(rows=1, cols=1)
        img2 = tab.rows[0].cells[0].add_paragraph()
        r = img2.add_run()
        r.add_picture(img, width=Cm(1.5), height=Cm(1.5))
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        table = document.add_table(rows=10, cols=18, style='Table Grid')

        a = table.cell(0, 0)
        b = table.cell(1, 14)
        A = a.merge(b)
        A.text = 'Renseignements donnés par :' + nomFr + ' ' + prenomFr + '  ' + 'Livret remis à :' + nomFr + ' ' + prenomFr
        A.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        A.left_indent = Inches(0.5)

        c = table.cell(0, 15)
        d = table.cell(1, 16)
        B = c.merge(d)
        B.text = 'Carnet n°: \n Page n°'
        B.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# -- PROVINCE --
        e = table.cell(2, 0)
        f = table.cell(2, 1)
        C = e.merge(f)
        C.text = 'Province de'
        C.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        g = table.cell(2, 2)
        h = table.cell(2, 4)
        D = g.merge(h)
        D.text = prov
        run = D.paragraphs[0].runs[0]
        run.font.name = 'Cambria'

        e = table.cell(2, 5)
        f = table.cell(2, 13)
        j = e.merge(f)
        j.text = 'RECONNAISSANCE PARCELLAIRE'
        run = j.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(15)
        j.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        j.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#Sous_zone
        h = table.cell(2, 14)
        i = table.cell(2, 17)
        k = h.merge(i)
        k.text = 'Sous-Secteur (Zone) n°:' + sZone
        k.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# -- Cercle  --
        c1 = table.cell(3, 0)
        c2 = table.cell(3, 1)
        c3 = c1.merge(c2)
        c3.text = 'Cercle de'
        c3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        c4 = table.cell(3, 2)
        c5 = table.cell(3, 4)
        c6 = c4.merge(c5)
        c6.text = cercle
        run = c6.paragraphs[0].runs[0]
        run.font.name = 'Cambria'

# -- Commune Rurale --
        c7 = table.cell(3, 5)
        c8 = table.cell(3, 13)
        c9 = c7.merge(c8)
        c9.text = 'Commune Rurale : ' + ' ' + commun
        c9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c9.paragraphs[0].runs[0]
        run.font.size = Pt(13)

        c10 = table.cell(3, 14)
        c11 = table.cell(3, 17)
        c12 = c10.merge(c11)
        c12.text = 'Date :' + ' ' + str(datetime.datetime.now())
        c12.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# -- NOM ET ADRESSE DES PROPRIETAIRES --
        x1 = table.cell(4, 0)
        x2 = table.cell(8, 4)
        x3 = x1.merge(x2)
        x3.text = 'NOM ET ADRESSE DES \n PROPRIETAIRES \n (1)'
        x3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        x3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        nom = table.cell(9, 0)
        nom2 = table.cell(9, 4)
        nom3 = nom.merge(nom2)
        nom3.text = nomFr + ' ' +prenomFr + '\n' + adress + '\n CIN' + ' ' + cin + '\n TEL' + ' ' + tel
        nom3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        nom3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = nom3.paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)

# --------------------------------- RENSEIGNEMENTS SUR LES PARCELLES --------------------
        x4 = table.cell(4, 5)
        x5 = table.cell(4, 17)
        x6 = x4.merge(x5)
        x6.text = 'RENSEIGNEMENTS SUR LES PARCELLES'
        x6.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = x6.paragraphs[0].runs[0]
        run.font.size = Pt(12)

# -- DOUAR --
        x7 = table.cell(5, 5)
        x8 = table.cell(7, 6)
        x9 = x7.merge(x8)
        x9.text = 'Douar'
        x9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        x9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        z10 = table.cell(8, 5)
        zx11 = table.cell(8, 6)
        zx12 = z10.merge(zx11)
        zx12.text = '(2)'
        z10.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z10.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        drVal = table.cell(9, 5)
        drVal2 = table.cell(9, 6)
        drVal3 = drVal.merge(drVal2)
        drVal3.text = douar
        drVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        drVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = drVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- NUM_PARCELLE --
        x10 = table.cell(5, 7)
        x11 = table.cell(7, 7)
        x12 = x10.merge(x11)
        x12.text = 'N° Plle \n / \n N° Réq'
        x12.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        x12.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z11 = table.cell(8, 7)
        z11.text = '(3)'
        z11.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z11.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        numPolyg = table.cell(9, 7)
        numPolyg.text = numParcel
        numPolyg.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        numPolyg.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = numPolyg.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- PARCELLE --
        y = table.cell(5, 8)
        y1 = table.cell(7, 10)
        y2 = y.merge(y1)
        y2.text = 'Nom \n de la parcelle'
        y2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z18 = table.cell(8, 8)
        z19 = table.cell(8, 10)
        z20 = z18.merge(z19)
        z20.text = '(4)'
        z20.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z20.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        nom_polyg = table.cell(9, 8)
        nom_polyg_2 = table.cell(9, 10)
        nom_polyg_val = nom_polyg.merge(nom_polyg_2)
        nom_polyg_val.text = nomParcel
        nom_polyg_val.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nom_polyg_val.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = nom_polyg_val.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- PROPRIETAIRE --
        y12 = table.cell(5, 11)
        y13 = table.cell(7, 14)
        y14 = y12.merge(y13)
        y14.text = 'Nom du propriétaire'
        y14.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y14.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z16 = table.cell(8, 11)
        zX16 = table.cell(8, 14)
        ZX16 = z16.merge(zX16)
        ZX16.text = '(5)'
        ZX16.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ZX16.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        propr = table.cell(9, 11)
        propr2 = table.cell(9, 14)
        propr3 = propr.merge(propr2)
        propr3.text = nomFr + ' ' + prenomFr
        propr3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        propr3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = propr3.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- Motif Opposition --
        z7 = table.cell(5, 15)
        z8 = table.cell(7, 17)
        z9 = z7.merge(z8)
        z9.text = "Motif de l'opposition"
        z9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        a1 = table.cell(8, 15)
        a2 = table.cell(8, 17)
        a3 = a1.merge(a2)
        a3.text = '(6)'
        a3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        a3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        motifOpp = table.cell(9, 15)
        motifOpp2 = table.cell(9, 17)
        motifOpp3 = motifOpp.merge(motifOpp2)
        motifOpp3.text = motifOppos
        motifOpp3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        motifOpp3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = motifOpp3.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- \\ INFORMATION BAS DE PAGE // --
        paragph1 = document.add_paragraph(" \n (1) A compléter par le numéro CINE ( La Carte Nationale d'Identité Électronique) et numéro de GSM le cas échéant")
        paragph1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph1.paragraph_format
        paragraph_format.left_indent = Inches(0.5)

        paragph2 = document.add_paragraph("(6) Consistance Matérielle : TC ( Terrain de Culture) –TP (Terrain Implanté) – TB (Terrain Bâti) \n A compléter en Arabe par :")
        paragph2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph2.paragraph_format
        paragraph_format.left_indent = Inches(0.5)

        table2 = document.add_table(rows=1, cols=12, style='Table Grid')
        table2.alignment = WD_TABLE_ALIGNMENT.RIGHT

        table2.cell(0, 0).text = 'فلاحيةأرض'
        table2.cell(0, 1).text = 'ارض بورية'
        table2.cell(0, 2).text = 'ارض زراعية مسقية'
        table2.cell(0, 3).text = 'ارض مغروسة'
        table2.cell(0, 4).text = 'ارض مسقية مغروسة'
        table2.cell(0, 5).text = 'ارض عارية'
        table2.cell(0, 6).text = 'أحراش'
        table2.cell(0, 7).text = 'أرض بها بناية'
        table2.cell(0, 8).text = 'ارض بها بنايات'
        table2.cell(0, 9).text = 'مقبرة'
        table2.cell(0, 10).text = 'مسجد'
        table2.cell(0, 11).text = 'مركب عقاري البيعة  -  مرفق رياضي'

        for column_index in range(len(table2.columns)):
            table2.cell(0, column_index).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table2.cell(0, column_index).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            column_index += 1

        for col_index in range(len(table2.columns)):
            if col_index == 11:
                table2.cell(0, col_index).width = Cm(3.5)
            else:
                table2.cell(0, col_index).width = Cm(2)
            col_index += 1

        paragph3 = document.add_paragraph("\n (7) Type de Spéculation : TR ( Culture Traditionnelle)- IN ( Culture Industrielle)- MA (Culture Maraichère-potagère)")
        paragph3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph3.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format = paragph3.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.line_spacing = Pt(6.5)

        paragph4 = document.add_paragraph("(8) Type de sol : RME (R’mel)- HAM (Hamri)- TIR (Tirs)- DEH (Dehs)- BIA (Beida).")
        paragph4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph4.paragraph_format.line_spacing = 1.0
        paragraph_format = paragph4.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.line_spacing = Pt(6.5)

        paragph5 = document.add_paragraph("N.B : Les carnets ZN1 doivent être établis sur calque stable ordinaire non végétal permettant leur conservation durable.")
        paragph5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph5.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.line_spacing = Pt(6.5)

        try:
            filename = 'P' + numParcel + '_ZN3.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3', filename)
            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3'))
            document.save(filepath)
            msg = messagebox.showinfo("ZN3", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            Zn3File.Zn3File()

        except:
            msg = messagebox.showerror("ZN1", "Veuillez fermer le fichier")


    def word_to_pdf(self, nomFr, prenomFr, prov, sZone, cercle, commun, adress, cin, tel, douar, numParcel, nomParcel, motifOppos):
        project_directory = os.path.abspath(os.curdir)
        img = project_directory + "/icons/ancfcc2.png"

        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        for section in document.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(297)  # for A4 paper
            section.page_height = Mm(210)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

# -- IMAGE --
        tab = document.add_table(rows=1, cols=1)
        img2 = tab.rows[0].cells[0].add_paragraph()
        r = img2.add_run()
        r.add_picture(img, width=Cm(1.5), height=Cm(1.5))
        img2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        img2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        table = document.add_table(rows=10, cols=18, style='Table Grid')

        a = table.cell(0, 0)
        b = table.cell(1, 14)
        A = a.merge(b)
        A.text = 'Renseignements donnés par :' + nomFr + ' ' + prenomFr + '  ' + 'Livret remis à :' + nomFr + ' ' + prenomFr
        A.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        A.left_indent = Inches(0.5)

        c = table.cell(0, 15)
        d = table.cell(1, 16)
        B = c.merge(d)
        B.text = 'Carnet n°: \n Page n°'
        B.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# -- PROVINCE --
        e = table.cell(2, 0)
        f = table.cell(2, 1)
        C = e.merge(f)
        C.text = 'Province de'
        C.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        g = table.cell(2, 2)
        h = table.cell(2, 4)
        D = g.merge(h)
        D.text = prov
        run = D.paragraphs[0].runs[0]
        run.font.name = 'Cambria'

        e = table.cell(2, 5)
        f = table.cell(2, 13)
        j = e.merge(f)
        j.text = 'RECONNAISSANCE PARCELLAIRE'
        run = j.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(15)
        j.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        j.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#Sous_zone
        h = table.cell(2, 14)
        i = table.cell(2, 17)
        k = h.merge(i)
        k.text = 'Sous-Secteur (Zone) n°:' + sZone
        k.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# -- Cercle  --
        c1 = table.cell(3, 0)
        c2 = table.cell(3, 1)
        c3 = c1.merge(c2)
        c3.text = 'Cercle de'
        c3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        c4 = table.cell(3, 2)
        c5 = table.cell(3, 4)
        c6 = c4.merge(c5)
        c6.text = cercle
        run = c6.paragraphs[0].runs[0]
        run.font.name = 'Cambria'

# -- Commune Rurale --
        c7 = table.cell(3, 5)
        c8 = table.cell(3, 13)
        c9 = c7.merge(c8)
        c9.text = 'Commune Rurale : ' + ' ' + commun
        c9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = c9.paragraphs[0].runs[0]
        run.font.size = Pt(13)

        c10 = table.cell(3, 14)
        c11 = table.cell(3, 17)
        c12 = c10.merge(c11)
        c12.text = 'Date :' + ' ' + str(datetime.datetime.now())
        c12.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# -- NOM ET ADRESSE DES PROPRIETAIRES --
        x1 = table.cell(4, 0)
        x2 = table.cell(8, 4)
        x3 = x1.merge(x2)
        x3.text = 'NOM ET ADRESSE DES \n PROPRIETAIRES \n (1)'
        x3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        x3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        nom = table.cell(9, 0)
        nom2 = table.cell(9, 4)
        nom3 = nom.merge(nom2)
        nom3.text = nomFr + ' ' +prenomFr + '\n' + adress + '\n CIN' + ' ' + cin + '\n TEL' + ' ' + tel
        nom3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        nom3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = nom3.paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(10)

# --------------------------------- RENSEIGNEMENTS SUR LES PARCELLES --------------------
        x4 = table.cell(4, 5)
        x5 = table.cell(4, 17)
        x6 = x4.merge(x5)
        x6.text = 'RENSEIGNEMENTS SUR LES PARCELLES'
        x6.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = x6.paragraphs[0].runs[0]
        run.font.size = Pt(12)

# -- DOUAR --
        x7 = table.cell(5, 5)
        x8 = table.cell(7, 6)
        x9 = x7.merge(x8)
        x9.text = 'Douar'
        x9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        x9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        z10 = table.cell(8, 5)
        zx11 = table.cell(8, 6)
        zx12 = z10.merge(zx11)
        zx12.text = '(2)'
        z10.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z10.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        drVal = table.cell(9, 5)
        drVal2 = table.cell(9, 6)
        drVal3 = drVal.merge(drVal2)
        drVal3.text = douar
        drVal.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        drVal.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = drVal.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- NUM_PARCELLE --
        x10 = table.cell(5, 7)
        x11 = table.cell(7, 7)
        x12 = x10.merge(x11)
        x12.text = 'N° Plle \n / \n N° Réq'
        x12.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        x12.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z11 = table.cell(8, 7)
        z11.text = '(3)'
        z11.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z11.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        numPolyg = table.cell(9, 7)
        numPolyg.text = numParcel
        numPolyg.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        numPolyg.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = numPolyg.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- PARCELLE --
        y = table.cell(5, 8)
        y1 = table.cell(7, 10)
        y2 = y.merge(y1)
        y2.text = 'Nom \n de la parcelle'
        y2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z18 = table.cell(8, 8)
        z19 = table.cell(8, 10)
        z20 = z18.merge(z19)
        z20.text = '(4)'
        z20.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z20.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        nom_polyg = table.cell(9, 8)
        nom_polyg_2 = table.cell(9, 10)
        nom_polyg_val = nom_polyg.merge(nom_polyg_2)
        nom_polyg_val.text = nomParcel
        nom_polyg_val.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        nom_polyg_val.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = nom_polyg_val.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- PROPRIETAIRE --
        y12 = table.cell(5, 11)
        y13 = table.cell(7, 14)
        y14 = y12.merge(y13)
        y14.text = 'Nom du propriétaire'
        y14.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        y14.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        z16 = table.cell(8, 11)
        zX16 = table.cell(8, 14)
        ZX16 = z16.merge(zX16)
        ZX16.text = '(5)'
        ZX16.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ZX16.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        propr = table.cell(9, 11)
        propr2 = table.cell(9, 14)
        propr3 = propr.merge(propr2)
        propr3.text = nomFr + ' ' + prenomFr
        propr3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        propr3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = propr3.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- Motif Opposition --
        z7 = table.cell(5, 15)
        z8 = table.cell(7, 17)
        z9 = z7.merge(z8)
        z9.text = "Motif de l'opposition"
        z9.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        z9.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        a1 = table.cell(8, 15)
        a2 = table.cell(8, 17)
        a3 = a1.merge(a2)
        a3.text = '(6)'
        a3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        a3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        motifOpp = table.cell(9, 15)
        motifOpp2 = table.cell(9, 17)
        motifOpp3 = motifOpp.merge(motifOpp2)
        motifOpp3.text = motifOppos
        motifOpp3.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        motifOpp3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = motifOpp3.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# -- \\ INFORMATION BAS DE PAGE // --
        paragph1 = document.add_paragraph(" \n (1) A compléter par le numéro CINE ( La Carte Nationale d'Identité Électronique) et numéro de GSM le cas échéant")
        paragph1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph1.paragraph_format
        paragraph_format.left_indent = Inches(0.5)

        paragph2 = document.add_paragraph("(6) Consistance Matérielle : TC ( Terrain de Culture) –TP (Terrain Implanté) – TB (Terrain Bâti) \n A compléter en Arabe par :")
        paragph2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph2.paragraph_format
        paragraph_format.left_indent = Inches(0.5)

        table2 = document.add_table(rows=1, cols=12, style='Table Grid')
        table2.alignment = WD_TABLE_ALIGNMENT.RIGHT

        table2.cell(0, 0).text = 'فلاحيةأرض'
        table2.cell(0, 1).text = 'ارض بورية'
        table2.cell(0, 2).text = 'ارض زراعية مسقية'
        table2.cell(0, 3).text = 'ارض مغروسة'
        table2.cell(0, 4).text = 'ارض مسقية مغروسة'
        table2.cell(0, 5).text = 'ارض عارية'
        table2.cell(0, 6).text = 'أحراش'
        table2.cell(0, 7).text = 'أرض بها بناية'
        table2.cell(0, 8).text = 'ارض بها بنايات'
        table2.cell(0, 9).text = 'مقبرة'
        table2.cell(0, 10).text = 'مسجد'
        table2.cell(0, 11).text = 'مركب عقاري البيعة  -  مرفق رياضي'

        for column_index in range(len(table2.columns)):
            table2.cell(0, column_index).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table2.cell(0, column_index).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            column_index += 1

        for col_index in range(len(table2.columns)):
            if col_index == 11:
                table2.cell(0, col_index).width = Cm(3.5)
            else:
                table2.cell(0, col_index).width = Cm(2)
            col_index += 1

        paragph3 = document.add_paragraph("\n (7) Type de Spéculation : TR ( Culture Traditionnelle)- IN ( Culture Industrielle)- MA (Culture Maraichère-potagère)")
        paragph3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph3.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format = paragph3.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.line_spacing = Pt(6.5)

        paragph4 = document.add_paragraph("(8) Type de sol : RME (R’mel)- HAM (Hamri)- TIR (Tirs)- DEH (Dehs)- BIA (Beida).")
        paragph4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragph4.paragraph_format.line_spacing = 1.0
        paragraph_format = paragph4.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.line_spacing = Pt(6.5)

        paragph5 = document.add_paragraph("N.B : Les carnets ZN1 doivent être établis sur calque stable ordinaire non végétal permettant leur conservation durable.")
        paragph5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format = paragph5.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.line_spacing = Pt(6.5)

        try:
            filename = 'P' + numParcel + '_ZN3.docx'
            filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3', filename)
            document.save(filepath)
            # PDF
            wordFile = 'P' + numParcel + '_ZN3.docx'
            pdfFile = 'P' + numParcel + '_ZN3.pdf'
            wdFormatPDF = 17
            in_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3', wordFile)
            out_file = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3', pdfFile)

            if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3')):
                os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3'))

            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=wdFormatPDF)
            doc.close()
            word.Quit()
            msg = messagebox.showinfo("ZN3", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
            Zn3File.Zn3File()

        except:
            msg = messagebox.showerror("ZN3", "Veuillez fermer le fichier")





