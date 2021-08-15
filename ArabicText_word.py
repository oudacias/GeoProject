import datetime
from tkinter import messagebox
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Mm
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.enum.table import WD_TABLE_DIRECTION
import os
import comtypes.client
from AdminAPP import Zn3File
from pyarabic.unshape import unshaping_word

from pyarabic import araby
def word():
    project_directory = os.path.abspath(os.curdir)
    img = project_directory + "/icons/ancfcc2.png"

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

    for section in document.sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Mm(297)  # for A4 paper
        section.page_height = Mm(210)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

# -- \\ INFORMATION BAS DE PAGE // --


    table2 = document.add_table(rows=1, cols=8, style='Table Grid')
    table2.alignment = WD_TABLE_ALIGNMENT.RIGHT

    table2.cell(0, 0).text = 'فلاحيةأرض'
    table2.cell(0, 1).text = u'ارض بورية'
    table2.cell(0, 2).text = u'ارض زراعية مسقية'
    table2.cell(0, 3).text = 'ارض مغروسة'
    table2.cell(0, 4).text = 'ارض مسقية مغروسة'
    table2.cell(0, 5).text = 'ارض عارية'
    table2.cell(0, 6).text = 'أحراش'
    table2.cell(0, 7).text = 'أرض بها بناية'
    table2.cell(0, 8).text = 'ارض بها بنايات'
    table2.cell(0, 9).text = 'مقبرة'
    table2.cell(0, 10).text = 'مسجد'
    table2.cell(0, 11).text = 'مركب عقاري البيعة  -  مرفق رياضي'


    for column_index in range(len(table2.columns)):
        table2.cell(0, column_index).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table2.cell(0, column_index).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        column_index += 1

    for col_index in range(len(table2.columns)):
        if col_index == 11:
            table2.cell(0, col_index).width = Cm(3.5)
        else:
            table2.cell(0, col_index).width = Cm(2)
        col_index += 1

    try:
        filename = 'Parcelle_ZN3.docx'
        filepath = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3', filename)
        if not os.path.exists(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3')):
            os.makedirs(os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop\IFE_PIECES\ZN3'))
        document.save(filepath)
        msg = messagebox.showinfo("ZN3", "Fichier bien généré, Pouvez-vous l'ouvrire maintenant")
        Zn3File.Zn3File()
        print(111)
    except:
        msg = messagebox.showerror("ZN3", "Veuillez fermer le fichier")

x = word()