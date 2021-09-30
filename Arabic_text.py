import docx
from docx.enum.style import WD_STYLE_TYPE



doc = docx.Document()
style = doc.styles['Normal']
font = style.font

font.rtl = True
doc.add_paragraph("ارض مسقية مغروسة")
doc.add_paragraph("HEllo World")
doc.save("path.docx")